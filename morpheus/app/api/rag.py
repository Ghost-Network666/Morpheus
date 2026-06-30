import os
import uuid
from fastapi import APIRouter, Depends, HTTPException, Request, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.database import get_db
from app.api.auth import require_user
from app.models.user import User
from app.models.document import Document
from app.core import rag_engine
from app.config import settings

router = APIRouter(prefix="/api/rag", tags=["rag"])


@router.post("/documents")
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(require_user),
):
    upload_dir = os.path.join(settings.data_dir, "uploads")
    os.makedirs(upload_dir, exist_ok=True)

    doc_id = str(uuid.uuid4())
    filename = file.filename or "document"
    save_path = os.path.join(upload_dir, f"{doc_id}_{filename}")

    content_bytes = await file.read()
    with open(save_path, "wb") as f:
        f.write(content_bytes)

    # Extract text
    text = _extract_text(filename, content_bytes)

    # Index
    chunks = await rag_engine.add_document(text, metadata={"filename": filename, "user_id": user.id}, doc_id=doc_id)

    doc = Document(
        id=doc_id,
        user_id=user.id,
        filename=filename,
        path=save_path,
        size=len(content_bytes),
        chunks=chunks,
    )
    db.add(doc)
    await db.commit()

    return _doc_out(doc)


@router.get("/documents")
async def list_documents(db: AsyncSession = Depends(get_db), user: User = Depends(require_user)):
    result = await db.execute(select(Document).where(Document.user_id == user.id))
    return [_doc_out(d) for d in result.scalars().all()]


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, db: AsyncSession = Depends(get_db), user: User = Depends(require_user)):
    result = await db.execute(select(Document).where(Document.id == doc_id, Document.user_id == user.id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    await rag_engine.delete_document(doc_id)
    if os.path.isfile(doc.path):
        os.remove(doc.path)
    await db.delete(doc)
    await db.commit()
    return {"ok": True}


@router.post("/query")
async def query_documents(request: Request, user: User = Depends(require_user)):
    body = await request.json()
    query = body.get("query", "")
    n = body.get("n_results", 5)
    if not query:
        raise HTTPException(400, "query required")

    results = await rag_engine.query(query, n_results=n, filter_metadata={"user_id": user.id})
    return {"results": results}


def _doc_out(d: Document) -> dict:
    return {
        "id": d.id,
        "filename": d.filename,
        "size": d.size,
        "chunks": d.chunks,
        "user_id": d.user_id,
        "created_at": d.created_at,
    }


def _extract_text(filename: str, content: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in (".docx",):
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        else:
            return content.decode("utf-8", errors="replace")
    except Exception:
        return content.decode("utf-8", errors="replace")
